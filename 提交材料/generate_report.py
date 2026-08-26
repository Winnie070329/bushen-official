# -*- coding: utf-8 -*-
"""Generate summer practice report docx."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE = Path(__file__).resolve().parent
PHOTOS = BASE / "照片原图"
OUT = BASE / "我用AI从零完成企业官网项目-实践报告.docx"

TITLE = "我用AI从零完成企业官网项目——以布神家居官网为例"
PRACTICE_TYPE = "AI智行"
CLASS_NAME = "【请填写班级】"
STUDENT_NAME = "【请填写姓名】"
DATE = "2026年8月"


def set_run_font(run, name="仿宋_GB2312", size=16, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = Pt(29)
    p.paragraph_format.space_after = Pt(6)
    if level == 1:
        run = p.add_run(text)
        set_run_font(run, "黑体", 16, True)
    elif level == 2:
        run = p.add_run(text)
        set_run_font(run, "楷体_GB2312", 16, True)
    else:
        run = p.add_run(text)
        set_run_font(run, "仿宋_GB2312", 16, True)
    return p


def add_body(doc, text, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = Pt(29)
    p.paragraph_format.space_after = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run, "仿宋_GB2312", 16, False)
    return p


def add_image(doc, path, caption):
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(5.8))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Cm(0)
    r = cap.add_run(caption)
    set_run_font(r, "仿宋_GB2312", 14, False)


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # Cover info
    for line in [
        f"题    目：{TITLE}",
        f"实践类型：{PRACTICE_TYPE}",
        f"班    级：{CLASS_NAME}",
        f"姓    名：{STUDENT_NAME}",
        "",
        DATE,
    ]:
        p = doc.add_paragraph(line)
        p.paragraph_format.line_spacing = Pt(29)
        for run in p.runs:
            set_run_font(run, "仿宋_GB2312", 16)

    doc.add_page_break()

    add_heading(doc, "目  录", 1)
    for item in [
        "一、实践背景",
        "二、时间地点",
        "三、主要内容",
        "四、取得成效",
        "五、总结体会",
    ]:
        add_body(doc, item, indent=False)

    doc.add_page_break()

    add_heading(doc, "一、实践背景", 1)
    add_body(
        doc,
        "在人工智能技术快速普及的背景下，AI 编程助手正在改变软件开发的学习与实践方式。"
        "计算机学院 2026 年暑期社会实践推出“AI智行”主题，鼓励同学们将 AI 工具应用于真实场景，"
        "完成从需求分析、方案设计到作品落地的完整过程。参考选题中的“我用AI从零完成××项目”"
        "全流程日志，正好契合我利用 AI 辅助开发企业官网的实践目标。",
    )
    add_body(
        doc,
        "本次实践选取“布神家居企业官网建设”作为项目载体。佛山市顺德区布神家居用品有限公司"
        "成立于 2012 年，主营窗帘布艺、纺织家纺与家居用品销售，位于顺德龙江家居产业带。"
        "通过公开信息检索，我发现该企业虽有基础网络信息，但缺乏一个结构清晰、视觉统一、"
        "可在移动端正常访问的展示型官网。借助 Cursor 等 AI 编程工具，我希望在较短时间内"
        "从零完成网站页面设计、代码编写、版本管理与线上部署，探索 AI 赋能中小企业数字化展示的"
        "可行路径，同时提升自身的前端开发、工程化协作与问题解决能力。",
    )

    add_heading(doc, "二、时间地点", 1)
    add_body(doc, "实践时间：2026 年 8 月 20 日至 2026 年 8 月 26 日。")
    add_body(doc, "实践地点：湖南省长沙市（居家线上实践）。")
    add_body(
        doc,
        "实践形式：以线上调研、AI 辅助编程、GitHub 代码托管与 GitHub Pages 静态部署为主，"
        "结合公开企业信息检索与网站效果自测，完成项目全流程。",
    )

    add_heading(doc, "三、主要内容", 1)
    add_heading(doc, "（一）需求调研与选题确定", 2)
    add_body(
        doc,
        "实践初期，我通过搜索引擎与公开工商信息，检索“佛山布神有限公司”“布神家居”等关键词，"
        "确认目标企业为佛山市顺德区布神家居用品有限公司，整理其成立时间、注册地址、经营范围等"
        "基础资料，并参考同类家纺企业官网的结构，确定网站需包含首页展示、关于我们、产品展示、"
        "业务范围与联系方式等模块。",
    )

    add_heading(doc, "（二）AI 辅助从零搭建网站", 2)
    add_body(
        doc,
        "在 Cursor 中，我向 AI 助手描述需求：“创建可在 GitHub Pages 展示的布神家居企业官网”。"
        "AI 首先生成了包含 HTML、CSS、JavaScript 的基础项目结构，实现了响应式导航、首页 Hero 区、"
        "关于我们、业务范围与联系信息等页面模块。随后我提出“搜索佛山布神有限公司并添加丰富图片”，"
        "AI 根据公开信息更新了企业简介与联系方式，并生成窗帘、床品、面料等主题氛围图，"
        "压缩优化后放入 images 目录，使网站视觉更贴近家纺行业特征。",
    )
    add_body(
        doc,
        "开发过程中，我并非简单复制 AI 输出，而是逐步理解页面结构：index.html 负责语义化布局，"
        "style.css 控制配色与栅格排版，main.js 实现移动端菜单交互。遇到 GitHub CLI 未登录、"
        "PowerShell 不支持 heredoc 等问题时，AI 会给出替代命令并最终完成 git init、commit 与 push。",
    )

    add_image(
        doc,
        PHOTOS / "01-AI辅助开发企业官网首页效果.jpg",
        "图1 AI 辅助开发完成的企业官网首页线上效果",
    )

    add_heading(doc, "（三）GitHub 仓库创建与 Pages 部署", 2)
    add_body(
        doc,
        "代码完成后，我将项目推送至 GitHub 仓库 bushen-official。在 AI 指导下，使用 gh auth login"
        "完成 GitHub CLI 登录后，通过 API 开启 GitHub Pages，将 main 分支根目录作为发布源。"
        "约两分钟后，网站成功上线，访问地址为 https://winnie070329.github.io/bushen-official/。"
        "这一步骤让我体会到：现代 Web 开发不仅是写页面，还包括版本管理、远程协作与自动化部署。",
    )

    add_image(
        doc,
        PHOTOS / "02-GitHub仓库与代码托管.jpg",
        "图2 GitHub 仓库代码托管与项目文件结构",
    )

    add_heading(doc, "（四）内容完善与效果验证", 2)
    add_body(
        doc,
        "网站第二版更新中，我补充了产品展示板块，将窗帘布艺、床品家纺、纺织面料三类核心业务"
        "以图文卡片形式呈现，并调整整体配色为暖色系家纺风格。完成后，我分别访问首页、产品区、"
        "关于我们等锚点链接，检查移动端菜单、图片加载与文字排版，确保在不同分辨率下均可正常浏览。",
    )

    add_image(
        doc,
        PHOTOS / "03-网站产品展示板块.jpg",
        "图3 网站产品展示板块效果截图",
    )
    add_image(
        doc,
        PHOTOS / "04-网站关于我们与联系信息板块.jpg",
        "图4 网站关于我们与联系信息板块效果截图",
    )

    add_heading(doc, "四、取得成效", 1)
    add_body(
        doc,
        "第一，完成了可公开访问的企业展示网站。项目包含 5 张优化后的产品氛围图、完整的页面模块"
        "与响应式布局，已成功部署至 GitHub Pages，实现了从“零代码”到“可访问 URL”的完整闭环。",
    )
    add_body(
        doc,
        "第二，掌握了 AI 辅助开发的基本流程。我学会了如何用自然语言向 AI 描述需求、如何分阶段"
        "迭代功能、如何在 AI 建议基础上理解 HTML/CSS/JS 结构，而不是被动接受结果。",
    )
    add_body(
        doc,
        "第三，提升了工程化实践能力。包括 Git 版本管理、GitHub 远程仓库推送、GitHub Pages 静态"
        "站点发布，以及使用命令行工具排查认证与部署问题。",
    )
    add_body(
        doc,
        "第四，将专业学习与社会实践结合。项目以真实企业公开信息为依据，探索 AI 在中小企业品牌"
        "展示、低成本建站场景中的应用价值，符合“AI智行”主题中“AI 文化传播与数字内容创作”"
        "的实践方向。",
    )

    add_heading(doc, "五、总结体会", 1)
    add_body(
        doc,
        "本次“我用AI从零完成企业官网项目”实践，让我深刻感受到 AI 并不是替代思考的工具，"
        "而是放大学习效率的助手。真正有价值的，是我在 AI 协作中主动完成的调研、需求拆解、"
        "效果验收与问题反馈。若缺少对目标企业信息的检索，网站内容就会空洞；若缺少对部署流程"
        "的理解，代码就无法真正上线服务访问者。",
    )
    add_body(
        doc,
        "实践也暴露了一些不足：例如部分图片为 AI 生成的示意氛围图，尚未替换为企业实拍素材；"
        "联系方式仍依赖公开信息，后续可进一步与企业核实。未来我希望继续学习组件化前端框架、"
        "可访问性规范与后端表单接口，使网站从“展示型”升级为“可交互、可维护”的完整产品。",
    )
    add_body(
        doc,
        "总体而言，这次暑期社会实践让我在真实项目中走完了“调研—开发—部署—总结”的全流程，"
        "更加理解了人工智能时代下计算机专业学生应具备的复合能力：既会用 AI，也懂基础原理；"
        "既能快速实现原型，也能对结果负责并不断改进。这对我今后的课程学习、竞赛实践与就业"
        "规划都具有积极意义。",
    )

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
